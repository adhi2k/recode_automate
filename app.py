import streamlit as st
import json
import base64
import re
import os
import time
import datetime
import qrcode
import docx
from docx.shared import Inches
from docx2pdf import convert
from github import Github, InputGitTreeElement

def load_courses():
    try:
        with open("courses.json", "r") as f:
            data = json.load(f)
            return data.get("courses", [])
    except Exception as e:
        st.error(f"Failed to load courses.json: {e}")
        return []

def load_config():
    try:
        with open("config.json", "r") as f:
            return json.load(f)
    except:
        return {}

def modify_readme(content, name, register_number):
    # Replace Name (use [ \t]* instead of \s* to avoid matching newlines and deleting the next line)
    content = re.sub(r'###[ \t]*Name:[ \t]*.*', f'### Name: {name}', content, flags=re.IGNORECASE)
    # Replace Register Number
    content = re.sub(r'###[ \t]*Register Number:[ \t]*.*', f'### Register Number: {register_number}', content, flags=re.IGNORECASE)
    
    # If the source repo didn't have a Register Number line at all, let's inject it under Name just in case
    if f'### Register Number: {register_number}' not in content:
        content = content.replace(f'### Name: {name}', f'### Name: {name}\n### Register Number: {register_number}')
        
    return content

def get_all_files(repo, branch="main"):
    contents = repo.get_contents("", ref=branch)
    files = []
    while contents:
        file_content = contents.pop(0)
        if file_content.type == "dir":
            contents.extend(repo.get_contents(file_content.path, ref=branch))
        else:
            files.append(file_content)
    return files

def automate_experiment_api(g, exp, name, register_number, progress_bar, status_text):
    # Backward compatibility with 'repo' key
    college_repo_url = exp.get("college_repo", exp.get("repo"))
    source_repo_url = exp.get("source_repo", college_repo_url)
    
    def extract_repo_name(url):
        if url and url.startswith("http"):
            parts = url.rstrip("/").split("/")
            return f"{parts[-2]}/{parts[-1]}"
        return url
        
    college_repo_name = extract_repo_name(college_repo_url)
    source_repo_name = extract_repo_name(source_repo_url)
    
    user = g.get_user()
    
    status_text.text(f"Preparing to fork from: {college_repo_name}...")
    progress_bar.progress(5)
    
    try:
        college_repo = g.get_repo(college_repo_name)
    except Exception as e:
        return False, f"Failed to access college repository {college_repo_name}: {e}"
        
    # Check if user already has a fork
    target_repo_basename = college_repo_name.split('/')[-1]
    try:
        target_repo = user.get_repo(target_repo_basename)
        status_text.text(f"Found existing repository: {target_repo.full_name}")
    except:
        status_text.text(f"Forking college repository {college_repo_name}...")
        try:
            target_repo = user.create_fork(college_repo)
            status_text.text("Waiting for GitHub to prepare the fork...")
            time.sleep(8) # Wait for fork to complete
            target_repo = user.get_repo(target_repo_basename)
        except Exception as e:
            return False, f"Failed to fork repository: {e}"
            
    progress_bar.progress(20)
    
    status_text.text(f"Fetching completed code from source: {source_repo_name}...")
    try:
        source_repo = g.get_repo(source_repo_name)
    except Exception as e:
        return False, f"Failed to access source repository {source_repo_name}: {e}"
        
    source_branch = source_repo.default_branch
    target_branch = target_repo.default_branch
    
    files = get_all_files(source_repo, branch=source_branch)
    progress_bar.progress(50)
    
    tree_elements = []
    
    status_text.text("Modifying details in memory...")
    for file_content in files:
        if file_content.name.lower() == "readme.md":
            raw_content = file_content.decoded_content.decode('utf-8')
            modified_content = modify_readme(
                raw_content, 
                name, 
                register_number
            )
            blob = target_repo.create_git_blob(modified_content, "utf-8")
            tree_elements.append(InputGitTreeElement(path=file_content.path, mode="100644", type="blob", sha=blob.sha))
        else:
            try:
                content = file_content.decoded_content
                encoded_content = base64.b64encode(content).decode('utf-8')
                blob = target_repo.create_git_blob(encoded_content, "base64")
                tree_elements.append(InputGitTreeElement(path=file_content.path, mode="100644", type="blob", sha=blob.sha))
            except Exception as e:
                pass
                
    progress_bar.progress(85)
    status_text.text("Committing changes directly to your GitHub...")
    
    try:
        ref = target_repo.get_git_ref(f"heads/{target_branch}")
        base_commit = target_repo.get_git_commit(ref.object.sha)
        base_tree = base_commit.tree
        
        new_tree = target_repo.create_git_tree(tree_elements, base_tree)
        new_commit = target_repo.create_git_commit("Automated upload of student details", new_tree, [base_commit])
        ref.edit(new_commit.sha)
    except Exception as e:
        return False, f"Failed to push commit: {e}", None
        
    progress_bar.progress(100)
    
    zip_url = f"{target_repo.html_url}/archive/refs/heads/{target_branch}.zip"
    
    success_message = f"Successfully pushed! \n\n" \
                      f"🔗 **[View Repository]({target_repo.html_url})**"
                      
    return True, success_message, target_repo.html_url

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000EE')
    rPr.append(c)

    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    
    t = docx.oxml.shared.OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def generate_record_document(name, register_number, selected_experiments, results, course_name):
    try:
        doc = docx.Document()
        
        # Add Logo if it exists
        if os.path.exists("image.png"):
            logo_p = doc.add_paragraph()
            logo_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_p.add_run()
            logo_run.add_picture("image.png", width=docx.shared.Inches(4.0))
        
        # Heading
        h1 = doc.add_paragraph()
        h1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run = h1.add_run("Table Of Contents")
        run.bold = True
        run.font.size = docx.shared.Pt(14)
        
        h2 = doc.add_paragraph()
        h2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        run = h2.add_run(course_name.upper())
        run.bold = True
        
        # Table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ex. No'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Title'
        hdr_cells[3].text = 'QR Code'
        hdr_cells[4].text = 'Marks'
        hdr_cells[5].text = 'Signature'
        
        temp_qrs = []
        for exp in selected_experiments:
            exp_id = exp['id']
            if exp_id in results:
                url = results[exp_id]
                row_cells = table.add_row().cells
                row_cells[0].text = str(exp_id)
                row_cells[1].text = "" # Date empty
                
                # Title cell
                title_text = f"{exp['title']}\nURL: "
                row_cells[2].text = title_text
                p = row_cells[2].paragraphs[0]
                add_hyperlink(p, url, url)
                
                # QR Code
                qr = qrcode.make(url)
                temp_qr_path = f"temp_qr_{exp_id}.png"
                qr.save(temp_qr_path)
                temp_qrs.append(temp_qr_path)
                
                p_qr = row_cells[3].paragraphs[0]
                run_qr = p_qr.add_run()
                run_qr.add_picture(temp_qr_path, width=docx.shared.Inches(1.0))
                
                row_cells[4].text = ""
                row_cells[5].text = ""
                
        # Signature block
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("I confirm that the GitHub links provided are entirely my own work.")
        doc.add_paragraph(f"NAME: {name}                                                                           REGISTER NUMBER: {register_number}")
        doc.add_paragraph("DATE:                                                                        LEARNER SIGNATURE:")
        
        doc.save("final_record.docx")
        
        # Clean up temp QR codes
        for qr_file in temp_qrs:
            if os.path.exists(qr_file):
                os.remove(qr_file)
                
        # Convert to PDF
        try:
            convert("final_record.docx", "final_record.pdf")
            return True, "Success"
        except Exception as e:
            return True, f"Word doc generated, but PDF conversion failed: {e}"
            
    except Exception as e:
        return False, f"Failed to generate document: {e}"
def main():
    st.set_page_config(page_title="GitHub Assignment Automator", page_icon="🎓", layout="wide")
    
    st.title("🎓 GitHub Assignment Automator")
    st.write("Easily fork and submit your college assignments using your GitHub Token.")
    
    config = load_config()
    courses = load_courses()
    
    with st.sidebar:
        st.header("Student Details")
        name = st.text_input("Name", value=config.get("name", ""))
        register_number = st.text_input("Register Number", value=config.get("register_number", ""))
        st.markdown("---")
        
    if not courses:
        st.warning("No courses found in courses.json")
        return
        
    course_names = [c["name"] for c in courses]
    selected_course_name = st.selectbox("Select Course", course_names)
    
    selected_course = next(c for c in courses if c["name"] == selected_course_name)
    experiments = selected_course["experiments"]
    
    st.subheader(f"Experiments for {selected_course_name}")
    
    # Select All Checkbox
    select_all = st.checkbox("Select All Experiments")
    
    selected_experiments = []
    
    cols = st.columns(3)
    for i, exp in enumerate(experiments):
        with cols[i % 3]:
            # Use select_all to control default state
            is_selected = st.checkbox(exp["title"], value=select_all, key=f"exp_{exp['id']}")
            if is_selected:
                selected_experiments.append(exp)
                
    st.markdown("---")
    
    tab1, tab2 = st.tabs(["🚀 GitHub Automation", "📝 Record Generator"])
    
    with tab1:
        st.subheader("Automate GitHub Fork & Push")
        token_default = config.get("github_token", "")
        if token_default == "YOUR_GITHUB_PERSONAL_ACCESS_TOKEN":
            token_default = ""
        token = st.text_input("GitHub Token", value=token_default, type="password", help="Requires 'repo' scope.")
    
        if st.button("🚀 Run GitHub Automation", type="primary", use_container_width=True):
            if not token:
                st.error("Please enter your GitHub Token.")
                st.stop()
            if not name or not register_number:
                st.error("Please enter your Name and Register Number in the sidebar.")
                st.stop()
                
            if not selected_experiments:
                st.warning("Please select at least one experiment.")
                st.stop()
                
            g = Github(token)
            try:
                user = g.get_user()
                st.success(f"Successfully authenticated as GitHub User: **{user.login}**")
            except Exception as e:
                st.error("Invalid GitHub Token. Authentication failed.")
                st.stop()
                
            st.info(f"Starting automation for {len(selected_experiments)} experiment(s)...")
            
            for exp in selected_experiments:
                st.markdown(f"### ⚙️ Processing: {exp['title']}")
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                success, msg, repo_url = automate_experiment_api(g, exp, name, register_number, progress_bar, status_text)
                
                if success:
                    status_text.success(msg)
                else:
                    status_text.error(msg)
                    
                time.sleep(1) # Small pause before next iteration
                
            st.success("🎉 GitHub Automation Complete!")

    with tab2:
        st.subheader("Generate Record Document")
        st.write("Generate the final Word/PDF document based on your GitHub username. (Does not require GitHub Token)")
        
        github_username = st.text_input("Your GitHub Username", placeholder="e.g. adhi2k")
        
        if st.button("📝 Generate Record", type="primary", use_container_width=True):
            if not github_username:
                st.error("Please enter your GitHub Username.")
                st.stop()
            if not name or not register_number:
                st.error("Please enter your Name and Register Number in the sidebar.")
                st.stop()
            if not selected_experiments:
                st.warning("Please select at least one experiment.")
                st.stop()
                
            st.info(f"Generating records for {len(selected_experiments)} experiment(s)...")
            
            results = {}
            for exp in selected_experiments:
                college_repo_url = exp.get("college_repo", exp.get("repo"))
                repo_basename = college_repo_url.rstrip("/").split("/")[-1]
                target_repo_url = f"https://github.com/{github_username}/{repo_basename}"
                results[exp['id']] = target_repo_url
                
            success, msg = generate_record_document(name, register_number, selected_experiments, results, selected_course_name)
            
            if success:
                st.success("🎉 Document Generation Complete! Your records are ready.")
                
                col1, col2 = st.columns(2)
                with col1:
                    if os.path.exists("final_record.pdf"):
                        with open("final_record.pdf", "rb") as pdf_file:
                            st.download_button(
                                label="📄 Download Record (PDF)",
                                data=pdf_file,
                                file_name=f"{register_number}_Record.pdf",
                                mime="application/pdf",
                                type="primary",
                                use_container_width=True
                            )
                with col2:
                    if os.path.exists("final_record.docx"):
                        with open("final_record.docx", "rb") as docx_file:
                            st.download_button(
                                label="📝 Download Record (Word)",
                                data=docx_file,
                                file_name=f"{register_number}_Record.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="secondary",
                                use_container_width=True
                            )
            else:
                st.error(msg)

if __name__ == "__main__":
    main()
